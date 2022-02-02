import webbrowser
from pathlib import Path
import docx
import string, random
import re
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx2pdf import convert
from docx.shared import Pt, RGBColor,Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# Provide the Folder path where you want to save the files
output = Path('H:\max') # Important

#Main Heading
heading_text = "[ONLINE] [TeamA] [TeamB] [HEADING_TAIL] [LIVESTREAM_WORDS]Live Stream NCAA Collage Football 02 September 2021 Full HD Coverage."
heading_tail_list = ['Live Stream', 'Game live', 'Op game', 'Live Reddit ', 'Live Free On Tv']
livestream_word_list = ['livestream', 'Live Hd', 'Add more here']
heading_color=RGBColor(0, 0, 0)

#Second Heading
# second_heading_text=''
second_heading_text = "Welcome To Watch [TeamA] vs. [TeamB]: [SECOND_HEADING_TAIL] [SECOND_LIVESTREAM_WORDS] How to live stream, TV channel, start time for Thursday's NCAA Football game. How to watch [TeamA] vs. [TeamB] football game"
second_heading_tail_list = ['Live Stream', 'Game live', 'Op game', 'Live Reddit ', 'Live Free On Tv']
second_livestream_word_list = ['livestream', 'Live Hd', 'Add more here']
second_heading_color=RGBColor(255, 5, 255)



PostInfo = {
    1: {
        'TeamA': "Germany vs",
        'TeamB': "Italy",
        'description': '''Enter Your description Here''',
        'Url_text':'Click Here TO Watch Live',
        'redirect_url':'https://www.youtube.com/',
        'file_format': '1video-[randomLD5]-[randomLD5]-[random1]-[random1]-[randomLD3]-[random1]-anything[random5]-[numbers].html',
        'image_location': {
            1: 'https://media.istockphoto.com/photos/close-up-of-legs-and-feet-of-football-player-in-blue-socks-and-shoes-picture-id1150952747?k=20&m=1150952747&s=612x612&w=0&h=vreccM0RO2rNp4aLN-mLyBwTfN7sfwvkdkwegzYPrXo=',
            2: 'https://encrypted-tbn0.gstatic.com/images?q=tbn:ANd9GcQIrvPUQyxbRZ0G4fv5NV_P2rB7fhyFf7Fdzg&usqp=CAU',
            3: 'G:\ProjectWorkHere\Projects_MK\#2-Docs2PDF\index3.jpg',
            4: 'https://i.imgur.com/IyC9C8m.jpeg',
            5: 'https://i.imgur.com/eQ2UNSI.jpeg',
            6: 'https://source.unsplash.com/random/?football',
            # 7: '#2-Docs2PDF\index.jpg',
            # 8: '#2-Docs2PDF\index.jpg',
            # 9: '#2-Docs2PDF\index.jpg',
            # 10: '#2-Docs2PDF\index.jpg',
        },
        
        #Make Description unique By adding those lines,otherwise leave it as it is..
        'FIRST_LINE':'Hello',
        'MIDDLE_LINE':'Welcome',
        'LAST_LINE':'Bye',
    },
    
     2: {
        'TeamA': "France v",
        'TeamB': "Norway",
        'description': '''Enter Your description Here''',
        'Url_text':'Click Here TO Watch Live',
        'redirect_url':'https://www.youtube.com/',
        'file_format': '1video-[randomLD5]-[randomLD5]-[random1]-[random1]-[randomLD3]-[random1]-anything[random5]-[numbers].html',
        'image_location': {
            1: 'https://media.istockphoto.com/photos/close-up-of-legs-and-feet-of-football-player-in-blue-socks-and-shoes-picture-id1150952747?k=20&m=1150952747&s=612x612&w=0&h=vreccM0RO2rNp4aLN-mLyBwTfN7sfwvkdkwegzYPrXo=',
            2: 'https://encrypted-tbn0.gstatic.com/images?q=tbn:ANd9GcQIrvPUQyxbRZ0G4fv5NV_P2rB7fhyFf7Fdzg&usqp=CAU',
            3: 'G:\ProjectWorkHere\Projects_MK\#2-Docs2PDF\index3.jpg',
            4: 'https://i.imgur.com/IyC9C8m.jpeg',
            5: 'https://i.imgur.com/eQ2UNSI.jpeg',
            6: 'https://source.unsplash.com/random/?football',
            # 7: '#2-Docs2PDF\index.jpg',
            # 8: '#2-Docs2PDF\index.jpg',
            # 9: '#2-Docs2PDF\index.jpg',
            # 10: '#2-Docs2PDF\index.jpg',
        },
        
        #Make Description unique By adding those lines,otherwise leave it as it is..
        'FIRST_LINE':'',
        'MIDDLE_LINE':'',
        'LAST_LINE':'',
    },
}


description_first_list = [
    'First line 1',
    'First line 2',
    'First line 3',
    'First line 4',
    'First line 5',
]

description_middle_list = [
    'MIDDLE line 1',
    'MIDDLE line 2',
    'MIDDLE line 3',
    'MIDDLE line 4',
    'MIDDLE line 5',
]

description_end_list = [
    'end line 1',
    'end line 2',
    'end line 3',
    'end line 4',
    'end line 5',
]


description_text = '''[FIRST_LINE]
Watch [TeamA] vs. [TeamB]: How to live stream, TV channel, start time for Thursday's NCAA Football game
How to watch [TeamA] vs. [TeamB] football game
Who's Playing
[TeamB] @ [TeamA]

Last Season Records: [TeamA] 0-6; [TeamB] 5-2

What to Know
The [TeamA] Rebels will play against a Division II opponent, the [TeamB] Eagles, in an early-season tune-up Thursday at 10 p.m. ET at Allegiant Stadium. Coming off of an uninspired 0-6 last-season record, the Rebels have set their aspirations higher this year.

The neutral point spread forecasts a close one for these two. We'll see if the game is as close as the oddsmakers expect or if one of these teams has a surprise blowout in them.

How To Watch
When: Thursday at 10 p.m. ET
Where: Allegiant Stadium -- Paradise, Nevada
TV: STADIUM
Follow: CBS Sports App
Ticket Cost: $15.00
Odds
The Rebels are a slight 2-point favorite against the Eagles, according to the latest college football odds.


[MIDDLE_LINE]



The oddsmakers had a good feel for the line for this one, as the game opened with the Rebels as a 3-point favorite.

Over/Under: -110

See college football picks for every single game, including this one, from SportsLine's advanced computer model. Get picks now.

Series History
This is the first time these teams have played each other within the last six years.

How to Watch [TeamB] Eagles  at [TeamA] Rebels: Live Stream, TV Channel, Start Time
[TeamA] looks to pick up its first win in two years when it hosts [TeamB] in the season opener Thursday night.

[TeamA] made a splash by signing transfer Tate Martell right before school started, but in a mild surprise, he was nowhere to be found when the Rebels released their depth chart. The Rebels, who have not won a game since 2019, need a spark and many thought he would be it. Instead, they will wait for him to play quarterback at least another week.

How to Watch:

Date: Sept. 2, 2021

Time: 10:00 p.m. ET

TV: Stadium

You can stream the game on fuboTV: Start with a 7-day free trial!

[TeamA] has a rough start to its schedule, and if it can't find a way to beat [TeamB], it could be a while before the Rebels win. After their opening game, they travel to No. 25 Arizona State before hosting Iowa State. They are going to be huge underdogs in those games, so they need to find a way to beat the Eagles.

[TeamB] went 5-2 in the spring of 2021 after the fall schedule was postponed due to COVID-19. The Eagles lost to North Dakota State in the first round of the playoffs after losing just one game in the regular season.

The Eagles are not going to be an easy game for the Rebels, and if you are looking for an FCS team to upset an FBS team this would be the game to watch. The Rebels haven't had a winning season since 2013 and have had a hard time winning even four games in a year. [TeamB] will give them everything they can handle and could easily win this game.

[LAST_LINE]

'''


output_directory = Path(output)
output_directory_pdf = output / 'pdf'
output_directory_docx = output / 'docx'


Path(output_directory).mkdir(parents=True, exist_ok=True)
Path(output_directory_docx).mkdir(parents=True, exist_ok=True)
Path(output_directory_pdf).mkdir(parents=True, exist_ok=True)


def docx_file_location(File_name) :
    return Path(output_directory_docx, File_name+'.docx')

def pdf_file_location(File_name) :
    return Path(output_directory_pdf, File_name+'.pdf')



def random_string(onlyString, no_of_string):
    if onlyString:
        return ''.join(random.choice(string.ascii_lowercase) for i in range(int(no_of_string)))
    return ''.join(random.choice(string.ascii_lowercase + string.digits) for i in range(int(no_of_string)))


def make_title(title, count):
    if '[numbers]' in title:
        if count < 10:
            title = title.replace('[numbers]', f'0{str(count)}')
        title = title.replace('[numbers]', str(count))
    if re.findall("random[0-9]", title):
        check_random = re.findall("random[0-9]", title)
        for num in range(len(check_random)):
            title = title.replace(f'[{check_random[num]}]', random_string(True, check_random[num][6]),1)
    if re.findall("randomLD[0-9]", title):
        check_randomLD = re.findall("randomLD[0-9]", title)
        for num in range(len(check_randomLD)):
            title = title.replace(f'[{check_randomLD[num]}]', random_string(False,check_randomLD[num][8]),1)

    return title
def make_description(description,post_no,teamA,teamB):
    if '[FIRST_LINE]' in description:
        if len(PostInfo[post_no]['FIRST_LINE'])==0:
            description = description.replace('[FIRST_LINE]',random.choice(description_first_list),1)
        else:
            description = description.replace('[FIRST_LINE]',PostInfo[post_no]['FIRST_LINE'],1)
            
    if '[LAST_LINE]' in description:
        if len(PostInfo[post_no]['LAST_LINE'])==0:
            description = description.replace('[LAST_LINE]',random.choice(description_end_list),1)
        else:
            description = description.replace('[LAST_LINE]',PostInfo[post_no]['LAST_LINE'],1)
    if '[MIDDLE_LINE]' in description:
        if len(PostInfo[post_no]['MIDDLE_LINE'])==0:
            description = description.replace('[MIDDLE_LINE]',random.choice(description_end_list),1)
        else:
            description = description.replace('[MIDDLE_LINE]',PostInfo[post_no]['MIDDLE_LINE'],1)
    if '[TeamA]' in description:
        description = description.replace('[TeamA]',teamA)
    if '[TeamB]' in description:
        description = description.replace('[TeamB]',teamB)
    return description



def make_heading(heading,teamA,teamB):
    if '[HEADING_TAIL]' in heading:
        heading = heading.replace('[HEADING_TAIL]',random.choice(heading_tail_list),1)
    if '[LIVESTREAM_WORDS]' in heading:
        heading = heading.replace('[LIVESTREAM_WORDS]',random.choice(livestream_word_list),1)
    if '[TeamA]' in heading:
        heading = heading.replace('[TeamA]',teamA)
    if '[TeamB]' in heading:
        heading = heading.replace('[TeamB]',teamB)
    return heading

def make_second_heading(heading,teamA,teamB):
    if '[SECOND_HEADING_TAIL]' in heading:
        heading = heading.replace('[SECOND_HEADING_TAIL]',random.choice(second_heading_tail_list),1)
    if '[SECOND_LIVESTREAM_WORDS]' in heading:
        heading = heading.replace('[SECOND_LIVESTREAM_WORDS]',random.choice(second_livestream_word_list),1)
    if '[TeamA]' in heading:
        heading = heading.replace('[TeamA]',teamA)
    if '[TeamB]' in heading:
        heading = heading.replace('[TeamB]',teamB)
    return heading










def writedocx(content, font_name = 'Times New Roman', font_size = 12, font_bold = False, font_italic = False, font_underline = False, color = RGBColor(0, 0, 0),
              before_spacing = 0, after_spacing = 0,
              align = 'left', style = ''):
    paragraph = doc.add_paragraph(str(content))
    paragraph.style = doc.styles.add_style(style+random_string(0,10), WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline
    font.color.rgb = color
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(before_spacing)
    paragraph_format.space_after = Pt(after_spacing)
    # paragraph.line_spacing = line_spacing
    # paragraph_format.keep_together = keep_together
    # paragraph_format.keep_with_next = keep_with_next
    # paragraph_format.page_break_before = page_break_before
    # paragraph_format.widow_control = widow_control
    if align.lower() == 'left':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif align.lower() == 'center':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif align.lower() == 'right':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif align.lower() == 'justify':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return paragraph




def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def image_hyperlink(image_location,redirected_url):
    new_paragraph = doc.add_paragraph()
    new_run = new_paragraph.add_run()
    redirected_to = new_run.add_picture(image_location, width=Inches(6.0), height=Inches(4.5))
    r_id = new_paragraph.part.relate_to(redirected_url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('a:hlinkClick')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    redirected_to._inline.docPr.append(hyperlink)

def imageFromWeb(url):
    from io import BytesIO

    import requests

    response = requests.get(url)  
    # https://c.ndtvimg.com/2020-05/5mk12sl_football-generic-afp_625x300_06_May_20.jpg
    # https://i.imgur.com/IyC9C8m.jpeg
    return BytesIO(response.content) 











# ------------------------------------
#Writing into to docs file from here
# ------------------------------------



count = 1

for post_no in range(1, len(PostInfo) + 1):
    TeamA = PostInfo[post_no]['TeamA']
    TeamB = PostInfo[post_no]['TeamB']
    TeamA_ = '-'.join(TeamA.split())
    TeamB_ = '-'.join(TeamB.split())
    for images in range(1, len(PostInfo[post_no]['image_location']) + 1):
        doc = Document()
        newdescription = make_description(description_text,post_no,TeamA,TeamB)
        newHeading = make_heading(heading_text,TeamA,TeamB)
        newSecondheading=make_second_heading(second_heading_text,TeamA,TeamB)
        
        heading = writedocx(newHeading,color=heading_color,font_bold=True,font_size=24,align='center',style='headStyle'+random_string(1,5),after_spacing=10,font_name='Cambria')
        after_heading = writedocx(newSecondheading,color=second_heading_color,font_bold=True,align='center',font_size=12,style='after_headStyle',after_spacing=20,font_name='Calibri')
        url = add_hyperlink(writedocx('',font_bold=True,align='center',font_size=15,after_spacing=10,style='url1'), PostInfo[post_no]['Url_text'], PostInfo[post_no]['redirect_url'])
        # binary_img = imageFromWeb('https://i.imgur.com/IyC9C8m.jpeg')
        image_link =PostInfo[post_no]['image_location'][images] 
        try:
            image_hyperlink(image_link,PostInfo[post_no]['redirect_url'])
        except:
            try:
                image_hyperlink(imageFromWeb(image_link),PostInfo[post_no]['redirect_url'])
            except:
                print(f"{image_link} not working ,Please download this image and try again")
                continue
        # image_hyperlink(PostInfo[post_no]['image_location'][images],PostInfo[post_no]['redirect_url'])
        
        
        
        
        url = add_hyperlink(writedocx('',font_bold=True,align='center',font_size=15,after_spacing=20), PostInfo[post_no]['Url_text'], PostInfo[post_no]['redirect_url'])

        description = writedocx(newdescription,font_bold=True,align='center',font_size=12,style='descStyle')

        File_name = make_title(PostInfo[post_no]['file_format'], count)

        # pdf_file_location(File_name)
        doc.save(docx_file_location(File_name))



        
        convert(docx_file_location(File_name), pdf_file_location(File_name))
        print(f"{count} Files successfully generated in '{output_directory}'")
        count = count+1

webbrowser.open(output_directory)  # Opens the output folder.


# new_paragraph = doc.add_paragraph()
# new_run = new_paragraph.add_run()
# facebook = new_run.add_picture(image_location, width=Inches(6.0))
# r_id = new_paragraph.part.relate_to('http://facebook.com', docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
# hyperlink = docx.oxml.shared.OxmlElement('a:hlinkClick')
# hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
# facebook._inline.docPr.append(hyperlink)